# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.

from docx import Document
def print_hi(name):
    # Use a breakpoint in the code line below to debug your script.
    print(f'Hi, {name}')  # Press Ctrl+F8 to toggle the breakpoint.





# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    print_hi('PyCharm')
    f = open('1.txt','r')
    data = f.readlines()
    # print(data,end='')
    for line in data:
        print(line, end="")

    print(data[0])
    print('一次性读取所有')
    data1 = f.read()
    print(data[0][2])
    print(data[1])
    print(data[2])
    f.close()
    # print(f.encoding)
    document = Document()
    document.add_heading('xyt',level=1)
    document.add_heading('xyt1',level=0)
    document.save('1.docx')
    document.add_paragraph('xxx')



# See PyCharm help at https://www.jetbrains.com/help/pycharm/
